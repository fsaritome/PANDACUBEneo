"""DOCX export (optional second output format alongside the searchable PDF).

Built from the pipeline's own ordered regions rather than PaddleX's built-in
`save_to_word`, because that would serialize PP-StructureV3's internal OCR
text while the PDF carries our configured engine's words - two outputs of the
same document disagreeing on their contents is not acceptable for documents
that get cited. Here both formats are rendered from one source of truth.

Structure is preserved where the layout model supplies it: recognized tables
become real Word tables, figure regions are cropped from the page raster and
embedded as images, and titles become headings. Page content is staged per
page during OCR (the plugin runs per page, often in another process) and
assembled once the sandwich finishes.
"""
from __future__ import annotations

import json
import logging
import uuid
from html.parser import HTMLParser
from pathlib import Path

from patent_ocr.layout.types import RegionKind
from patent_ocr.qc import page_sort_key, resolve_staging_dir

log = logging.getLogger(__name__)

_CONTENT_SUFFIX = ".page.json"

_BODY_KINDS = {RegionKind.COLUMN.value, RegionKind.FULL_PAGE.value}


class _TableHtmlParser(HTMLParser):
    """Minimal <table> reader: PaddleX emits plain table markup, and pulling in
    an HTML library just for it would be disproportionate."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.rows: list[list[str]] = []
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag, attrs):
        if tag == "tr":
            self._row = []
        elif tag in ("td", "th"):
            self._cell = []

    def handle_endtag(self, tag):
        if tag == "tr" and self._row is not None:
            if self._row:
                self.rows.append(self._row)
            self._row = None
        elif tag in ("td", "th") and self._cell is not None:
            if self._row is None:
                self._row = []
            self._row.append(" ".join("".join(self._cell).split()))
            self._cell = None

    def handle_data(self, data):
        if self._cell is not None:
            self._cell.append(data)


def parse_table_html(html: str) -> list[list[str]]:
    """Rectangular cell grid from table markup; [] when nothing usable."""
    parser = _TableHtmlParser()
    try:
        parser.feed(html)
        parser.close()
    except Exception:  # noqa: BLE001 - malformed OCR markup must not kill the export
        log.warning("could not parse table html; falling back to text")
        return []
    rows = [r for r in parser.rows if any(c.strip() for c in r)]
    if not rows:
        return []
    width = max(len(r) for r in rows)
    return [r + [""] * (width - len(r)) for r in rows]


def write_page_content(regions, source_name: str = "", page_image=None,
                       qc_dir: str | Path | None = None) -> None:
    """Stage one page's ordered blocks (and any figure crops) for assembly."""
    path = resolve_staging_dir(qc_dir)
    if path is None:
        return
    path.mkdir(parents=True, exist_ok=True)
    prefix = page_sort_key(source_name)

    blocks: list[dict] = []
    for index, region in enumerate(regions):
        kind = region.kind.value
        text = " ".join(w.text for w in region.words).strip()
        block: dict = {"kind": kind, "text": text}

        if region.kind == RegionKind.TABLE and region.html:
            block["html"] = region.html
        elif region.kind == RegionKind.FIGURE and page_image is not None:
            crop = _save_crop(page_image, region.bbox, path, f"{prefix}_{index}")
            if crop:
                block["image"] = crop
        if text or block.get("html") or block.get("image"):
            blocks.append(block)

    if not blocks:
        return
    name = f"{prefix}_{uuid.uuid4().hex}{_CONTENT_SUFFIX}"
    (path / name).write_text(json.dumps({"blocks": blocks}), encoding="utf-8")


def _save_crop(page_image, bbox, out_dir: Path, stem: str) -> str | None:
    """Write a figure region out as PNG so the DOCX can embed the real artwork."""
    try:
        from PIL import Image

        x0, y0, x1, y1 = (int(v) for v in bbox)
        if x1 - x0 < 8 or y1 - y0 < 8:  # slivers are layout noise, not artwork
            return None
        crop = Image.fromarray(page_image[y0:y1, x0:x1])
        target = out_dir / f"{stem}.fig.png"
        crop.save(target)
        return target.name
    except Exception:  # noqa: BLE001 - a bad crop must not fail the page
        log.warning("could not crop figure region %s", bbox, exc_info=True)
        return None


def read_pages(work_dir: Path) -> list[list[dict]]:
    """All staged pages, in page order (filenames carry the sort prefix)."""
    if not work_dir.exists():
        return []
    pages = []
    for file in sorted(work_dir.glob(f"*{_CONTENT_SUFFIX}")):
        try:
            pages.append(json.loads(file.read_text(encoding="utf-8"))["blocks"])
        except (OSError, ValueError, KeyError):
            log.warning("skipping unreadable page content file: %s", file)
    return pages


def write_docx(work_dir: Path, output_path: Path, title: str = "",
               strip_line_numbers: bool = False) -> bool:
    """Assemble staged pages into one .docx. Returns False if nothing was written."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        log.warning("python-docx not installed; skipping DOCX output (pip install python-docx)")
        return False

    pages = read_pages(work_dir)
    if not pages:
        return False

    document = Document()
    if title:
        document.add_heading(title, level=1)

    for index, blocks in enumerate(pages):
        if index:
            document.add_page_break()
        for block in blocks:
            kind = block.get("kind", "")

            if strip_line_numbers and kind == RegionKind.MARGIN_NUMBERS.value:
                continue

            if kind == RegionKind.TABLE.value and block.get("html"):
                if _add_table(document, block["html"]):
                    continue

            if kind == RegionKind.FIGURE.value and block.get("image"):
                image = work_dir / block["image"]
                if image.exists():
                    try:
                        document.add_picture(str(image), width=Inches(5.5))
                        continue
                    except Exception:  # noqa: BLE001 - unusable image falls through to text
                        log.warning("could not embed figure %s", image)

            if not block.get("text"):
                continue
            if kind == RegionKind.TITLE.value:
                document.add_heading(block["text"], level=2)
                continue

            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            if kind == RegionKind.MARGIN_NUMBERS.value:
                # Patent line numbers: kept, but visually subordinate so they
                # cannot be mistaken for claim text.
                run.font.size = Pt(8)
                run.italic = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif kind not in _BODY_KINDS:
                run.font.size = Pt(9)
                run.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return True


def _add_table(document, html: str) -> bool:
    grid = parse_table_html(html)
    if not grid:
        return False
    table = document.add_table(rows=len(grid), cols=len(grid[0]))
    table.style = "Table Grid"
    for r, row in enumerate(grid):
        for c, cell in enumerate(row):
            table.cell(r, c).text = cell
    return True

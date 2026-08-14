import json

import numpy as np
import pytest

from patent_ocr.docx_export import parse_table_html, write_docx, write_page_content
from patent_ocr.layout.types import Region, RegionKind
from patent_ocr.ocr.base import Word


def _region(kind: RegionKind, *words: str, bbox=(0, 0, 40, 40), html=None) -> Region:
    region = Region(kind=kind, bbox=bbox, order_index=0, html=html)
    region.words = [Word(w, (0, 0, 5, 5), 99.0, "e") for w in words]
    return region


def test_table_html_parses_into_a_rectangular_grid():
    html = "<table><tr><td>A</td><td>B</td></tr><tr><td>1</td><td>2</td></tr></table>"
    assert parse_table_html(html) == [["A", "B"], ["1", "2"]]


def test_ragged_rows_are_padded_to_equal_width():
    html = "<table><tr><td>A</td><td>B</td><td>C</td></tr><tr><td>1</td></tr></table>"
    assert parse_table_html(html) == [["A", "B", "C"], ["1", "", ""]]


def test_header_cells_and_nested_markup_are_handled():
    html = "<table><tr><th>H</th></tr><tr><td><b>x</b> y</td></tr></table>"
    assert parse_table_html(html) == [["H"], ["x y"]]


def test_unusable_table_html_yields_no_grid():
    assert parse_table_html("") == []
    assert parse_table_html("<table></table>") == []


def test_table_and_figure_blocks_are_staged(tmp_path, monkeypatch):
    monkeypatch.setenv("PATENT_OCR_QC_DIR", str(tmp_path))
    page = np.zeros((80, 80, 3), dtype=np.uint8)
    regions = [
        _region(RegionKind.TITLE, "CLAIMS"),
        _region(RegionKind.TABLE, html="<table><tr><td>A</td></tr></table>"),
        _region(RegionKind.FIGURE, bbox=(0, 0, 40, 40)),
    ]
    write_page_content(regions, "000001_rasterize.png", page)

    staged = json.loads(next(tmp_path.glob("*.page.json")).read_text())["blocks"]
    kinds = [b["kind"] for b in staged]
    assert "table" in kinds and "figure" in kinds and "title" in kinds
    figure = next(b for b in staged if b["kind"] == "figure")
    assert (tmp_path / figure["image"]).exists()


def test_sliver_regions_are_not_emitted_as_figures(tmp_path, monkeypatch):
    """A 4px-wide 'figure' is layout noise, not artwork."""
    monkeypatch.setenv("PATENT_OCR_QC_DIR", str(tmp_path))
    page = np.zeros((80, 80, 3), dtype=np.uint8)
    write_page_content([_region(RegionKind.FIGURE, bbox=(0, 0, 4, 80))],
                       "000001_rasterize.png", page)
    assert list(tmp_path.glob("*.page.json")) == []


def test_docx_contains_a_real_table_and_heading(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    (tmp_path / "000001_a.page.json").write_text(json.dumps({"blocks": [
        {"kind": "title", "text": "CLAIMS"},
        {"kind": "table", "text": "", "html":
            "<table><tr><td>A</td><td>B</td></tr><tr><td>1</td><td>2</td></tr></table>"},
        {"kind": "column", "text": "body text"},
    ]}), encoding="utf-8")

    out = tmp_path / "out.docx"
    assert write_docx(tmp_path, out) is True

    doc = Document(str(out))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(1, 1).text == "2"
    assert any(p.style.name.startswith("Heading") and p.text == "CLAIMS"
               for p in doc.paragraphs)


def test_malformed_table_html_falls_back_without_raising(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    (tmp_path / "000001_a.page.json").write_text(json.dumps({"blocks": [
        {"kind": "table", "text": "fallback text", "html": "<table><tr><td>"},
    ]}), encoding="utf-8")
    out = tmp_path / "out.docx"
    assert write_docx(tmp_path, out) is True
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "fallback text" in text


def _claims_page(tmp_path):
    (tmp_path / "000001_a.page.json").write_text(json.dumps({"blocks": [
        {"kind": "margin_numbers", "text": "5 10 15 20 25 30 35"},
        {"kind": "column", "text": "A spinal bone fastener assembly"},
    ]}), encoding="utf-8")


def test_line_numbers_are_kept_by_default(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    _claims_page(tmp_path)
    out = tmp_path / "keep.docx"
    write_docx(tmp_path, out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "5 10 15 20 25 30 35" in text
    assert "A spinal bone fastener assembly" in text


def test_line_numbers_can_be_stripped(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    _claims_page(tmp_path)
    out = tmp_path / "stripped.docx"
    write_docx(tmp_path, out, strip_line_numbers=True)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "5 10 15 20 25 30 35" not in text
    # Body text must survive untouched.
    assert "A spinal bone fastener assembly" in text

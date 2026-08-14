"""Report what structure survived into a DOCX."""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

for path in sorted(Path("harder_tests/result").glob("*.docx")):
    doc = Document(str(path))
    paras = [p for p in doc.paragraphs if p.text.strip()]
    headings = [p.text for p in paras if p.style.name.startswith("Heading")]
    images = doc.inline_shapes
    print(f"\n=== {path.name} ===")
    print(f"  paragraphs : {len(paras)}")
    print(f"  headings   : {len(headings)}")
    print(f"  tables     : {len(doc.tables)}")
    print(f"  images     : {len(images)}")
    for h in headings[:6]:
        print(f"    heading: {h[:70].encode('ascii','replace').decode()}")
    for t in doc.tables[:2]:
        print(f"    table {len(t.rows)}x{len(t.columns)}: "
              f"{[c.text[:18] for c in t.rows[0].cells][:5]}")
    for s in list(images)[:4]:
        print(f"    image: {s.width.inches:.1f}in x {s.height.inches:.1f}in")
